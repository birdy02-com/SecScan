# -*- coding: utf-8 -*-
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess, os, json, time, argparse, sys, asyncio, aiodns, re
from urllib.parse import urlparse
from docx.shared import RGBColor
from colorama import init, Fore
from datetime import datetime
from docx import Document

if sys.platform == 'win32': asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
init()
outPath = "output"


def program_exit():
    print("Program exited...")
    sys.exit(1)


class DNS:
    nameservers = [
        '114.114.114.114', '114.114.115.115', '114.114.114.119', '114.114.115.119', '114.114.114.110',
        '114.114.115.110',
        '223.5.5.5', '223.6.6.6',
        '180.76.76.76',
        '119.29.29.29', '119.28.28.28', '182.254.116.116', '182.254.118.118', '1.12.12.12', '120.53.53.53',
        '1.2.4.8', '210.2.4.8',
        '117.50.11.11', '52.80.66.66', '117.50.10.10', '52.80.52.52', '117.50.60.30', '52.80.60.30',
        '101.6.6.6',
        '202.141.162.123', '202.141.176.93', '202.38.93.153'
    ]
    nameservers.extend([
        '112.4.0.55', '112.4.0.55', '112.4.1.36', '112.4.1.36', '112.4.12.200', '112.4.12.200', '120.196.122.69',
        '120.196.122.69', '120.196.141.86', '120.196.141.86', '120.196.165.7', '120.196.165.7', '203.142.100.18',
        '203.142.100.18', '203.142.100.21', '203.142.100.21', '211.136.17.97', '211.136.17.98', '211.136.17.107',
        '211.136.17.107', '211.136.112.50', '211.136.112.50', '211.136.150.66', '211.136.150.66', '211.136.192.6',
        '211.136.192.6', '211.137.32.178', '211.137.32.178', '211.137.130.2', '211.137.130.3', '211.137.130.19',
        '211.137.130.19', '211.137.160.5', '211.137.160.5', '211.137.160.185', '211.137.160.185', '211.137.241.34',
        '211.137.241.34', '211.138.75.123', '211.138.75.123', '211.138.91.1', '211.138.91.2', '211.138.106.19',
        '211.138.106.19', '211.138.180.2', '211.138.180.3', '211.138.200.69', '211.138.200.69', '211.138.240.100',
        '211.138.240.100', '211.138.245.1', '211.138.245.1', '211.138.245.180', '211.138.245.180', '211.139.5.29',
        '211.139.5.30', '211.139.73.34', '211.139.73.34', '211.140.13.188', '211.140.13.188', '211.140.188.188',
        '211.140.188.188', '211.140.197.58', '211.140.197.58', '211.141.16.99', '211.141.16.99', '211.141.90.68',
        '211.141.90.68', '218.201.17.2', '218.201.17.2', '218.202.152.130', '218.202.152.131', '218.203.160.195',
        '218.203.160.195', '221.130.252.200', '221.130.252.200', '221.131.143.69', '221.131.143.69', '61.132.163.68',
        '202.102.213.68', '219.141.136.10', '219.141.140.10', '61.128.192.68', '61.128.128.68', '218.85.152.99',
        '218.85.157.99', '202.100.64.68', '61.178.0.93', '202.96.128.86', '202.96.128.166', '202.96.134.33',
        '202.96.128.68', '202.103.225.68', '202.103.224.68', '202.98.192.67', '202.98.198.167', '222.88.88.88',
        '222.85.85.85', '219.147.198.230', '219.147.198.242', '202.103.24.68', '202.103.0.68', '222.246.129.80',
        '59.51.78.211', '218.2.2.2', '218.4.4.4', '61.147.37.1', '218.2.135.1', '202.101.224.69', '202.101.226.68',
        '219.148.162.31', '222.74.39.50', '219.146.1.66', '219.147.1.66', '218.30.19.40', '61.134.1.4',
        '202.96.209.133', '116.228.111.118', '202.96.209.5', '108.168.255.118', '61.139.2.69', '218.6.200.139',
        '219.150.32.132', '219.146.0.132', '222.172.200.68', '61.166.150.123', '202.101.172.35', '61.153.177.196',
        '61.153.81.75', '60.191.244.5', '123.123.123.123', '123.123.123.124', '202.106.0.20', '202.106.195.68',
        '221.5.203.98', '221.7.92.98', '210.21.196.6', '221.5.88.88', '202.99.160.68', '202.99.166.4',
        '202.102.224.68', '202.102.227.68', '202.97.224.69', '202.97.224.68', '202.98.0.68', '202.98.5.68',
        '221.6.4.66', '221.6.4.67', '202.99.224.68', '202.99.224.8', '202.102.128.68', '202.102.152.3',
        '202.102.134.68', '202.102.154.3', '202.99.192.66', '202.99.192.68', '221.11.1.67', '221.11.1.68',
        '210.22.70.3', '210.22.84.3', '119.6.6.6', '124.161.87.155', '202.99.104.68', '202.99.96.68', '221.12.1.227',
        '221.12.33.227', '202.96.69.38', '202.96.64.68'])
    nameservers.extend([
        '37.235.1.174', '37.235.1.177', '216.146.35.35', '216.146.36.36', '8.26.56.26', '8.20.247.20', '64.6.64.6',
        '64.6.65.6', '77.88.8.8', '77.88.8.1', '80.80.80.80', '80.80.81.81', '94.140.14.14', '94.140.15.15',
        '94.140.14.15', '94.140.15.16', '156.154.70.1', '156.154.71.1', '156.154.70.5', '156.154.71.5',
        '103.197.104.178', '103.197.106.75', '203.189.136.148', '203.112.2.4', '84.200.69.80', '84.200.70.40',
        '195.46.39.39', '195.46.39.40', '109.69.8.51', '91.239.100.100', '89.233.43.71', '81.218.119.11',
        '209.88.198.133', '185.222.222.222', '45.11.45.11', '74.82.42.42', '66.220.18.42', '104.236.210.29',
        '45.55.155.25', '185.228.168.9', '185.228.169.9', '185.228.168.10', '185.228.169.11', '185.228.168.168',
        '185.228.169.168', '202.79.32.33', '202.79.32.34'])

    @classmethod
    def main(cls, dom) -> dict | None:
        data = asyncio.run(cls.async_run(dom))
        try:
            return {"domain": data.get("name"), "ip": data.get("ipv4")[0], "dnsServer": "", "timer": Function.timer()}
        except:
            return None

    @classmethod
    async def async_run(cls, domain: str) -> dict:
        resolver = aiodns.DNSResolver(nameservers=cls.nameservers)
        return await cls.resolve_dns(resolver, domain)

    @classmethod
    async def resolve_dns(cls, resolver, hostname: str, slf: bool = False) -> dict:
        try:
            result = await resolver.query(hostname, 'A')
            ipv4 = [i.host for i in result]
            return {'name': hostname, 'ipv4': ipv4}
        except Exception as e:
            if re.search("Timeout while contacting DNS servers", str(e)) and not slf:
                return await cls.resolve_dns(resolver, hostname, True)
            if re.search("Could not contact DNS servers", str(e)) and not slf:
                return await cls.resolve_dns(resolver, hostname, True)
            return {}


# 方法
class Function:
    dom_suffix = [
        '.ac.cn', '.ah.cn', '.bj.cn', '.com.cn', '.cq.cn', '.fj.cn', '.gd.cn', '.gov.cn', '.gs.cn', '.gx.cn', '.gz.cn',
        '.ha.cn',
        '.hb.cn', '.he.cn', '.hi.cn', '.hk.cn', '.hl.cn', '.hn.cn', '.jl.cn', '.js.cn', '.jx.cn', '.ln.cn', '.mo.cn',
        '.net.cn',
        '.nm.cn', '.nx.cn', '.org.cn', '.zj.cn', '.edu.cn', '.cn', '.com', '.edu', '.gov', '.net', '.org', '.biz',
        '.info', '.pro',
        '.name', '.museum', '.coop', '.aero', '.xxx', '.idv', '.xyz', '.asia', '.co', '.top', '.icu', '.site', '.cc',
        '.vip', '.tv',
        '.ltd', '.club', '.me', '.cfd', '.cloud', '.online', '.work', '.fun', '.cx', '.cm', '.pub', '.life', '.us',
        '.fr', '.games',
        '.link', '.in', '.tech', '.market', '.uk', '.live', '.tw', '.pw', '.ink', '.fit', '.shop', '.guru', '.store',
        '.website',
        '.wiki', '.cyou', '.pl', '.moe', '.mobi', '.hk', '.city', '.men', '.wang', '.bond', '.tokyo', '.one', '.hu',
        '.chat', '.host',
        '.so', '.space', '.cf', '.buzz', '.win', '.gq', '.bid', '.trade', '.loan', '.gdn', '.tel', '.date', '.vc',
        '.racing',
        '.science', '.ws', '.dev', '.la', '.nl', '.de', '.ne.jp', '.mil.jp', '.go.jp', '.ac.jp', '.or.jp', '.co.jp',
        '.jp', '.ga',
        '.ru', '.tk', '.bz', '.today', '.fi', '.co.cz', '.ml', '.ml', '.app', '.art', '.click', '.sbs'
    ]

    @classmethod
    def timer(cls) -> str:  # 时间
        now = datetime.now()
        year = f'{now.year:02}'
        month = f'{now.month:02}'
        day = f'{now.day:02}'
        hour = f'{now.hour:02}'
        minute = f'{now.minute:02}'
        second = f'{now.second:02}'
        return f"{year}年{month}月{day}日{hour}时{minute}分{second}秒"

    @classmethod
    def fileGetUrl(cls, file) -> list:
        try:
            with open(file, "r", encoding='utf-8') as f:
                return [i.strip() for i in f.read().split('\n') if i.strip() and i.startswith('http')]
        except Exception as e:
            print(e)
            return []

    @classmethod
    def fileGetLine(cls, file) -> list:
        try:
            with open(file, "r", encoding='utf-8') as f:
                return [i.strip() for i in f.read().split('\n') if i.strip()]
        except Exception as e:
            print(e)
            return []

    @classmethod
    def getRootDomain(cls, domain: str) -> str:
        if domain.startswith('http'): domain = urlparse(domain).hostname
        if len(domain.split('.')) > 1:
            for i in cls.dom_suffix:
                if domain.endswith(i):
                    domain = domain[:-len(i)].split('.')[-1] + i
                    return domain
        return ''


# 输出日志
def log(model: str = "", text: str = "", end: bool = False):
    out_text = f"\r[ {model} ] {Fore.RESET + text}" + Fore.RESET
    if end:
        print(end=Fore.GREEN + out_text)
    else:
        print(Fore.GREEN + out_text)


# 文件输出
class Docx:
    @classmethod
    def _row_tit(cls, doc, text: str):
        line = doc.add_paragraph()
        tit = line.add_run(f'{text}：')
        tit.bold = True
        tit.font.color.rgb = RGBColor(0, 51, 102)  # 使用RGB颜色
        return line

    @classmethod
    def write_docx(cls, v: dict, _url: str, out_file: str) -> str:
        if not os.path.exists(out_file): os.makedirs(out_file)
        doc = Document()
        # 标题
        heading = doc.add_heading('漏洞报告', 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t = doc.add_paragraph('生成时间：' + Function.timer())
        t.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        cls._row_tit(doc, "漏洞名称").add_run(v['vName'])
        cls._row_tit(doc, "漏洞URL").add_run(_url)
        cls._row_tit(doc, "是否存在漏洞").add_run('是' if v['isVul'] else '否')
        cls._row_tit(doc, "影响产品").add_run(v['product'])
        cls._row_tit(doc, "影响版本").add_run(v['version'])
        cls._row_tit(doc, "危险等级").add_run(v['level'])
        if v['vId']:
            cls._row_tit(doc, "漏洞编号")
            for n, i in enumerate(v['vId']):
                doc.add_paragraph(str(n + 1) + '：' + i)

        cls._row_tit(doc, "漏洞描述").add_run(v['vDesc'])
        cls._row_tit(doc, "参考链接").add_run(v['link'])
        if v['fix']:
            cls._row_tit(doc, "修复建议")
            for n, i in enumerate(v['fix']):
                doc.add_paragraph(str(n + 1) + '：' + i)
        doc.add_paragraph('')
        doc.add_paragraph('')
        request = v['request']
        cls._row_tit(doc, ">> 请求")
        cls._row_tit(doc, "URL").add_run(request['url'])
        cls._row_tit(doc, "请求方法").add_run(request['method'])
        cls._row_tit(doc, "请求头")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                hed, text = request['header'], ''
                for i in hed: text += i + ': ' + hed[i] + '\n'
                cell.text = text.rstrip('\n')
        if request['body']:
            cls._row_tit(doc, "请求体")
            table = doc.add_table(rows=1, cols=1)
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells): cell.text = request['body']
        doc.add_paragraph('')
        doc.add_paragraph('')
        response = v['response']
        cls._row_tit(doc, ">> 响应")
        cls._row_tit(doc, "URL").add_run(response['url'])
        cls._row_tit(doc, "状态码").add_run(str(response['code']))
        cls._row_tit(doc, "响应头")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                hed, text = response['header'], ''
                for i in hed: text += i + ': ' + hed[i] + '\n'
                cell.text = text.rstrip('\n')
        cls._row_tit(doc, "响应体")
        table = doc.add_table(rows=1, cols=1)
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells): cell.text = response['body']
        f_name = urlparse(request['url']).hostname + "_" + v['vName'].replace("/", '') + '.docx'
        doc.save(os.path.join(out_file, f_name))
        return f_name


# 漏洞检测方法
def exploit(_poc, _url: str) -> dict | None:
    try:
        process = subprocess.Popen(
            [
                "secScript.exe",
                "-poc", _poc,
                "-vUrl", _url,
                "-api", "true"
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8'  # 指定正确的编码
        )
        stdout, _ = process.communicate()
        return json.loads(stdout)
    except KeyboardInterrupt:
        program_exit()
    except:
        return None


# icp查询
# 返回参考：https://apifox.com/apidoc/shared-76f1bd0e-6083-4251-91a2-e96c9bb3bce2/api-219597580
def icp(keyword: str) -> dict | None:
    try:
        process = subprocess.Popen(
            [
                "secScript.exe",
                "-icp", keyword,
                "-api", "true"
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8'  # 指定正确的编码
        )
        stdout, _ = process.communicate()
        return json.loads(stdout)
    except KeyboardInterrupt:
        program_exit()
    except:
        return None


# ip属地查询
# 返回参考：https://apifox.com/apidoc/shared-76f1bd0e-6083-4251-91a2-e96c9bb3bce2/api-217943200
def ip(ipv4: str) -> dict | None:
    try:
        process = subprocess.Popen(
            [
                "secScript.exe",
                "-ip", ipv4,
                "-api", "true"
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8'  # 指定正确的编码
        )
        stdout, _ = process.communicate()
        return json.loads(stdout)
    except KeyboardInterrupt:
        program_exit()
    except:
        return None


# 获取域名的A记录
# 返回参考：https://apifox.com/apidoc/shared-76f1bd0e-6083-4251-91a2-e96c9bb3bce2/api-219607563
def dns(domain: str) -> dict | None:
    try:
        process = subprocess.Popen(
            [
                "secScript.exe",
                "-dns", domain,
                "-api", "true"
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8'  # 指定正确的编码
        )
        stdout, _ = process.communicate()
        return json.loads(stdout)
    except KeyboardInterrupt:
        program_exit()
    except:
        return DNS.main(domain)


# 分析URL信息
def analyze_url(uri: str, cms: bool = False) -> dict | None:
    try:
        process = subprocess.Popen(
            [
                "secScript.exe",
                "-url", uri,
                "-cms", "true" if cms else "",
                "-api", "true"
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8'  # 指定正确的编码
        )
        stdout, _ = process.communicate()
        return json.loads(stdout)
    except KeyboardInterrupt:
        program_exit()
    except:
        return None


# 获取POC基本信息
# json 返回本地poc的json概述
# list 返回产品的poc列表
def poc_info(typer: str) -> dict | None:
    try:
        if typer not in ['json', 'list']: return {}
        process = subprocess.Popen(
            [
                "secScript.exe",
                "-pocs", typer,
                "-api", "true",
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            encoding='utf-8'  # 指定正确的编码
        )
        stdout, _ = process.communicate()
        return json.loads(stdout)
    except KeyboardInterrupt:
        program_exit()
    except:
        return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('-poc', '--poc', help='poc文件')
    parser.add_argument('-pocs', '--pocs', help='本地poc的基本信息，可选值：json、list')
    parser.add_argument('-u', '--url', help='配合漏洞检测时候（要检测的URL） | 站点分析时（要分析的url）')
    parser.add_argument('-cms', '--cms', help='配合站点分析时使用，输入值不为空')
    parser.add_argument('-out', '--out', type=bool, default=False, help='是否输出文件：(True | False)')
    parser.add_argument('-ip', '--ipv4', help='要查询属地的ipv4地址')
    parser.add_argument('-icp', '--icp', help='要查询ICP的单位名/域名/ICP号')
    parser.add_argument('-uf', '--ufile', help='指定要批量检测的url文件，配合-poc使用')
    parser.add_argument('-ns', '--dns', help='要获取A记录的域名')
    args = parser.parse_args()

    if args.poc:
        path = os.path.join(outPath, str(int(time.time())))
        if not os.path.isfile(args.poc): return print("文件不存在", args.poc)
        if args.url:
            if not args.url.startswith('http'): return print("URL错误", args.url)
            res = exploit(args.poc, args.url)
            if args.out:
                print(args.url, "存在漏洞" if res["isVul"] else "不存在漏洞", res["vName"])
                if res['isVul']: print("输出文件:", Docx.write_docx(res, args.url, path))
            else:
                print(res)

        elif args.ufile:
            if not os.path.isfile(args.ufile): return print("文件不存在", args.ufile)
            with open(args.ufile, 'r', encoding='utf-8') as f:
                urls = [i.strip() for i in f.read().split('\n') if i.strip() != "" and i.startswith('http')]
            for url in urls:
                res = exploit(args.poc, url)
                print(url, "存在漏洞" if res["isVul"] else "不存在漏洞", res["vName"])
                if res['isVul']: print("输出文件:", Docx.write_docx(res, args.url, path))

    elif args.ipv4:
        res = ip(args.ipv4)
        print(res)

    elif args.icp:
        res = icp(args.icp)
        print(res)

    elif args.dns:
        res = dns(args.dns)
        print(res)
    elif args.url:
        res = analyze_url(args.url, args.cms)
        print(res)
    elif args.pocs:
        res = poc_info(args.pocs)
        print(res)
        # for i in res:
        #     print(i)


if __name__ == "__main__":
    main()

# poc = r"C:\Users\b0bef\Documents\Golang\secScript\EXE\vul\OA\致远OA\致远OA A6 createMysql.jsp 数据库敏感信息泄露-1.yaml"
# url = "https://api.birdy02.com"
